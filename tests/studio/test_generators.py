from pathlib import Path

from cleanpaste_studio import verifier
from cleanpaste_studio.generators import (
    generate_docx,
    generate_markdown,
    generate_pdf,
    generate_text,
)
from cleanpaste_studio.parser import parse_document
from cleanpaste_studio.templates import get_profile
from docx import Document
from pypdf import PdfReader


def sample_document():
    return parse_document(
        "# Heading\n\nText with [link](https://example.com).\n\n- Item\n\n| A | B |\n| - | - |\n| 1 | 2 |\n\n```py\nprint('ok')\n```",
        "Generated Fixture",
        "markdown",
    )


def test_all_generators_create_portable_files(tmp_path: Path):
    parsed = sample_document()
    profile = get_profile("standard_business_brief")
    docx = generate_docx(parsed, profile, tmp_path / "fixture.docx")
    pdf = generate_pdf(parsed, profile, tmp_path / "fixture.pdf")
    markdown = generate_markdown(parsed, tmp_path / "fixture.md")
    text = generate_text(parsed, tmp_path / "fixture.txt")

    assert docx.size > 10_000
    assert pdf.size > 1_000
    assert markdown.size > 80
    assert text.size > 50
    word = Document(docx.path)
    assert word.core_properties.title == "Generated Fixture"
    assert len(word.tables) == 1
    assert "Heading" in " ".join(paragraph.text for paragraph in word.paragraphs)
    assert len(PdfReader(pdf.path).pages) == 1
    pdf_text = PdfReader(pdf.path).pages[0].extract_text()
    assert pdf_text.count("Generated Fixture") == 1
    assert "&quot;" not in pdf_text
    assert "print('ok')" in pdf_text
    assert "[link](https://example.com)" in markdown.path.read_text(encoding="utf-8")
    assert "A | B" in text.path.read_text(encoding="utf-8")


def test_pdf_verifier_renders_a_non_blank_preview(tmp_path: Path):
    parsed = sample_document()
    pdf = generate_pdf(parsed, get_profile("google_docs_default"), tmp_path / "fixture.pdf")
    result = verifier.verify_pdf(pdf.path, tmp_path)
    assert result.page_count == 1
    if verifier._find_pdftoppm():
        assert result.status == "verified"
        assert result.preview_path and result.preview_path.is_file()
    else:
        assert result.status == "unverified"
        assert result.reason and "Poppler" in result.reason


def test_pdf_verifier_is_truthful_without_a_renderer(tmp_path: Path, monkeypatch):
    parsed = sample_document()
    pdf = generate_pdf(parsed, get_profile("google_docs_default"), tmp_path / "fixture.pdf")
    monkeypatch.setattr(verifier, "_find_pdftoppm", lambda: None)
    result = verifier.verify_pdf(pdf.path, tmp_path)
    assert result.status == "unverified"
    assert result.page_count == 1
    assert result.reason and "Poppler" in result.reason


def test_unknown_template_fails_closed():
    try:
        get_profile("not-a-template")
    except ValueError as exc:
        assert "Unknown template" in str(exc)
    else:
        raise AssertionError("Unknown template unexpectedly accepted")
