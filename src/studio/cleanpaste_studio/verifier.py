from __future__ import annotations

import os
import shutil
import subprocess
import sys
import threading
import zipfile
from pathlib import Path

from docx import Document
from PIL import Image, ImageStat
from pypdf import PdfReader

from .models import VerificationResult

_WORD_RENDER_LOCK = threading.Lock()


def _find_pdftoppm() -> str | None:
    explicit = os.environ.get("CLEAN_PASTE_PDFTOPPM")
    if explicit and Path(explicit).is_file():
        return explicit
    bundled = (
        Path.home()
        / ".cache/codex-runtimes/codex-primary-runtime/dependencies/native/poppler/Library/bin/pdftoppm.exe"
    )
    if bundled.is_file():
        return str(bundled)
    found = shutil.which("pdftoppm")
    return found


def verify_pdf(path: Path, preview_dir: Path | None = None) -> VerificationResult:
    checks: list[str] = []
    try:
        reader = PdfReader(str(path))
        page_count = len(reader.pages)
        if page_count < 1:
            return VerificationResult(path.name, "failed", 0, checks, "PDF contains no pages")
        checks.append(f"PDF opened successfully with {page_count} page(s)")
        extracted = "".join((page.extract_text() or "") for page in reader.pages[:3]).strip()
        if not extracted:
            return VerificationResult(
                path.name, "failed", page_count, checks, "Rendered document has no extractable text"
            )
        checks.append("Text extraction returned non-empty content")

        renderer = _find_pdftoppm()
        if not renderer:
            return VerificationResult(
                path.name,
                "unverified",
                page_count,
                checks,
                "Poppler renderer is unavailable; structural checks passed but visual rendering did not run",
            )
        preview_dir = preview_dir or path.parent
        preview_dir.mkdir(parents=True, exist_ok=True)
        prefix = preview_dir / f"{path.stem}-preview"
        subprocess.run(
            [renderer, "-f", "1", "-singlefile", "-png", "-r", "96", str(path), str(prefix)],
            check=True,
            capture_output=True,
            timeout=60,
        )
        preview = prefix.with_suffix(".png")
        with Image.open(preview) as image:
            grayscale = image.convert("L")
            stats = ImageStat.Stat(grayscale)
            if image.width < 400 or image.height < 500:
                return VerificationResult(
                    path.name,
                    "failed",
                    page_count,
                    checks,
                    "Rendered page dimensions are implausibly small",
                )
            if stats.var[0] < 0.5:
                return VerificationResult(
                    path.name, "failed", page_count, checks, "Rendered first page appears blank"
                )
        checks.append("First page rendered to PNG and passed non-blank image checks")
        return VerificationResult(path.name, "verified", page_count, checks, preview_path=preview)
    except Exception as exc:  # noqa: BLE001 - verifier must fail closed with a structured result
        return VerificationResult(path.name, "failed", None, checks, f"{type(exc).__name__}: {exc}")


def verify_docx(path: Path, preview_dir: Path | None = None) -> VerificationResult:
    checks: list[str] = []
    try:
        if not zipfile.is_zipfile(path):
            return VerificationResult(
                path.name, "failed", None, checks, "DOCX is not a valid OPC ZIP package"
            )
        checks.append("DOCX is a valid OPC ZIP package")
        document = Document(path)
        text = "".join(paragraph.text for paragraph in document.paragraphs).strip()
        table_text = "".join(
            cell.text for table in document.tables for row in table.rows for cell in row.cells
        ).strip()
        if not text and not table_text:
            return VerificationResult(
                path.name, "failed", None, checks, "DOCX contains no readable content"
            )
        checks.append("WordprocessingML parsed with non-empty document content")
        checks.append(
            f"Structure contains {len(document.paragraphs)} paragraph(s) and {len(document.tables)} table(s)"
        )
        if os.name != "nt":
            return VerificationResult(
                path.name,
                "unverified",
                None,
                checks,
                "Structural verification passed; visual DOCX rendering requires Microsoft Word on Windows",
            )
        preview_dir = preview_dir or path.parent
        preview_dir.mkdir(parents=True, exist_ok=True)
        rendered_pdf = preview_dir / f"{path.stem}-docx-render.pdf"
        pid_file = preview_dir / f"{path.stem}-word-render.pid"
        with _WORD_RENDER_LOCK:
            try:
                result = subprocess.run(
                    [
                        sys.executable,
                        "-m",
                        "cleanpaste_studio.word_render_helper",
                        str(path.resolve()),
                        str(rendered_pdf.resolve()),
                        str(pid_file.resolve()),
                    ],
                    capture_output=True,
                    text=True,
                    timeout=120,
                    check=False,
                )
            except subprocess.TimeoutExpired:
                _terminate_word_renderer(pid_file)
                return VerificationResult(
                    path.name,
                    "unverified",
                    None,
                    checks,
                    "Structural verification passed; Microsoft Word rendering exceeded 120 seconds",
                )
            finally:
                pid_file.unlink(missing_ok=True)
        if result.returncode != 0 or not rendered_pdf.is_file():
            reason = (
                result.stderr or result.stdout or "Microsoft Word renderer was unavailable"
            ).strip()
            return VerificationResult(
                path.name,
                "unverified",
                None,
                checks,
                f"Structural verification passed; visual rendering did not complete: {reason[:300]}",
            )
        checks.append("Microsoft Word opened the DOCX and exported a rendered PDF")
        rendered = verify_pdf(rendered_pdf, preview_dir)
        if rendered.status != "verified":
            return VerificationResult(
                path.name,
                rendered.status,
                rendered.page_count,
                checks + rendered.checks,
                rendered.reason,
            )
        checks.extend(rendered.checks)
        return VerificationResult(
            path.name,
            "verified",
            rendered.page_count,
            checks,
            preview_path=rendered.preview_path,
        )
    except Exception as exc:  # noqa: BLE001 - verifier must fail closed with a structured result
        return VerificationResult(path.name, "failed", None, checks, f"{type(exc).__name__}: {exc}")


def _terminate_word_renderer(pid_file: Path) -> None:
    """Terminate only the dedicated DispatchEx Word process recorded by our helper."""
    try:
        pid = int(pid_file.read_text(encoding="ascii").strip())
        if pid > 0:
            subprocess.run(
                ["taskkill.exe", "/PID", str(pid), "/T", "/F"],
                capture_output=True,
                timeout=15,
                check=False,
            )
    except (OSError, ValueError, subprocess.SubprocessError):
        return
