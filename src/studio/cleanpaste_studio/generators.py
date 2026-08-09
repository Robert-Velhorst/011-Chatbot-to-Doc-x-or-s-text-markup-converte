from __future__ import annotations

import hashlib
import html
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    HRFlowable,
    KeepTogether,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from .models import Block, GeneratedArtifact, InlineSpan, ParsedDocument
from .templates import TemplateProfile


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _artifact(format_name: str, path: Path) -> GeneratedArtifact:
    return GeneratedArtifact(format_name, path, sha256_file(path), path.stat().st_size)


def _display_blocks(parsed: ParsedDocument) -> list[Block]:
    """Avoid repeating a source H1 when it is identical to the explicit document title."""
    blocks = list(parsed.blocks)
    if blocks and blocks[0].kind == "heading" and blocks[0].level == 1:
        source_title = " ".join(blocks[0].text.casefold().split())
        document_title = " ".join(parsed.title.casefold().split())
        if source_title == document_title:
            return blocks[1:]
    return blocks


def _rgb(value: str) -> RGBColor:
    value = value.lstrip("#")
    return RGBColor(*(int(value[i : i + 2], 16) for i in (0, 2, 4)))


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _add_hyperlink(paragraph, text: str, url: str, color: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color.lstrip("#"))
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.extend([c, underline])
    run.append(props)
    text_element = OxmlElement("w:t")
    text_element.set(qn("xml:space"), "preserve")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_docx_spans(paragraph, spans: list[InlineSpan], profile: TemplateProfile) -> None:
    for span in spans:
        if span.href:
            _add_hyperlink(paragraph, span.text, span.href, profile.accent)
            continue
        parts = span.text.split("\n")
        for index, part in enumerate(parts):
            if index:
                paragraph.add_run().add_break(WD_BREAK.LINE)
            run = paragraph.add_run(part)
            run.bold = span.bold
            run.italic = span.italic
            if span.code:
                run.font.name = profile.mono_font
                run.font.size = Pt(profile.body_size - 0.5)
                run.font.color.rgb = _rgb("#A52133")


def _configure_docx_styles(document: Document, profile: TemplateProfile) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = profile.body_font
    normal.font.size = Pt(profile.body_size)
    normal.font.color.rgb = _rgb(profile.text)
    normal.paragraph_format.space_after = Pt(6 if profile.compact else 8)
    normal.paragraph_format.line_spacing = profile.line_spacing

    for level in range(1, 7):
        style = styles[f"Heading {level}"]
        style.font.name = profile.heading_font
        style.font.bold = True
        style.font.color.rgb = _rgb(profile.text if level > 1 else profile.accent)
        style.font.size = Pt({1: 24, 2: 17, 3: 13.5, 4: 11.5, 5: 10.5, 6: 10}[level])
        style.paragraph_format.space_before = Pt({1: 18, 2: 15, 3: 12, 4: 10, 5: 8, 6: 6}[level])
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = profile.mono_font
    code.font.size = Pt(max(8.5, profile.body_size - 1))
    code.font.color.rgb = _rgb("#172033")
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(6)
    code.paragraph_format.space_after = Pt(8)

    if "Quote Clean" not in styles:
        quote = styles.add_style("Quote Clean", WD_STYLE_TYPE.PARAGRAPH)
    else:
        quote = styles["Quote Clean"]
    quote.font.name = profile.body_font
    quote.font.italic = True
    quote.font.color.rgb = _rgb(profile.muted)
    quote.paragraph_format.left_indent = Inches(0.25)
    quote.paragraph_format.space_after = Pt(8)


def generate_docx(
    parsed: ParsedDocument, profile: TemplateProfile, path: Path
) -> GeneratedArtifact:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(profile.margin_inches)
    section.left_margin = section.right_margin = Inches(profile.margin_inches)
    _configure_docx_styles(document, profile)
    document.core_properties.title = parsed.title
    document.core_properties.subject = "Generated locally by AI Clean Paste Document Studio"
    document.core_properties.keywords = "local-first, document conversion"

    title = document.add_paragraph(style="Title")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = title.add_run(parsed.title)
    run.font.name = profile.heading_font
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = _rgb(profile.text)
    title.paragraph_format.space_after = Pt(16)

    for block in _display_blocks(parsed):
        if block.kind == "heading":
            paragraph = document.add_paragraph(style=f"Heading {min(max(block.level, 1), 6)}")
            _add_docx_spans(paragraph, block.spans, profile)
        elif block.kind == "paragraph":
            paragraph = document.add_paragraph()
            _add_docx_spans(paragraph, block.spans, profile)
        elif block.kind == "list_item":
            style = "List Number" if block.ordered else "List Bullet"
            paragraph = document.add_paragraph(style=style)
            paragraph.paragraph_format.left_indent = Inches(0.25 + 0.25 * block.level)
            paragraph.paragraph_format.first_line_indent = Inches(-0.18)
            paragraph.paragraph_format.space_after = Pt(3)
            _add_docx_spans(paragraph, block.spans, profile)
        elif block.kind == "blockquote":
            paragraph = document.add_paragraph(style="Quote Clean")
            _add_docx_spans(paragraph, block.spans, profile)
        elif block.kind == "code":
            paragraph = document.add_paragraph(style="Code Block")
            paragraph.paragraph_format.keep_together = True
            run = paragraph.add_run(block.text)
            run.font.name = profile.mono_font
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "F4F6FA")
            paragraph._p.get_or_add_pPr().append(shading)
        elif block.kind == "horizontal_rule":
            paragraph = document.add_paragraph()
            border = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:color"), profile.accent.lstrip("#"))
            border.append(bottom)
            paragraph._p.get_or_add_pPr().append(border)
        elif block.kind == "table" and block.rows:
            columns = max(len(row) for row in block.rows)
            table = document.add_table(rows=len(block.rows), cols=columns)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            _set_repeat_table_header(table.rows[0])
            for row_index, values in enumerate(block.rows):
                for column_index in range(columns):
                    cell = table.cell(row_index, column_index)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    text = values[column_index] if column_index < len(values) else ""
                    cell.text = text
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_after = Pt(2)
                        for run in paragraph.runs:
                            run.font.name = profile.body_font
                            run.font.size = Pt(profile.body_size - 0.5)
                            run.bold = row_index == 0
                    if row_index == 0:
                        shading = OxmlElement("w:shd")
                        shading.set(qn("w:fill"), "EAF0FB")
                        cell._tc.get_or_add_tcPr().append(shading)
            document.add_paragraph().paragraph_format.space_after = Pt(0)

    document.save(path)
    return _artifact("docx", path)


def _pdf_spans(spans: list[InlineSpan]) -> str:
    parts: list[str] = []
    for span in spans:
        value = html.escape(span.text).replace("\n", "<br/>")
        if span.code:
            value = f'<font name="Courier" color="#A52133">{value}</font>'
        if span.bold:
            value = f"<b>{value}</b>"
        if span.italic:
            value = f"<i>{value}</i>"
        if span.href:
            value = f'<a href="{html.escape(span.href, quote=True)}" color="#1456D9">{value}</a>'
        parts.append(value)
    return "".join(parts)


def generate_pdf(parsed: ParsedDocument, profile: TemplateProfile, path: Path) -> GeneratedArtifact:
    path.parent.mkdir(parents=True, exist_ok=True)
    page_size = A4 if profile.page_size == "a4" else LETTER
    document = SimpleDocTemplate(
        str(path),
        pagesize=page_size,
        rightMargin=profile.margin_inches * inch,
        leftMargin=profile.margin_inches * inch,
        topMargin=profile.margin_inches * inch,
        bottomMargin=profile.margin_inches * inch,
        title=parsed.title,
        author="AI Clean Paste Document Studio",
    )
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "CleanBody",
        parent=styles["BodyText"],
        fontName="Helvetica" if profile.body_font != "Georgia" else "Times-Roman",
        fontSize=profile.body_size,
        leading=profile.body_size * profile.line_spacing,
        textColor=colors.HexColor(profile.text),
        spaceAfter=6 if profile.compact else 8,
    )
    heading_styles = {
        level: ParagraphStyle(
            f"CleanHeading{level}",
            parent=styles[f"Heading{min(level, 4)}"],
            fontName="Helvetica-Bold",
            fontSize={1: 22, 2: 16, 3: 13, 4: 11, 5: 10, 6: 9.5}[level],
            leading={1: 26, 2: 20, 3: 16, 4: 14, 5: 13, 6: 12}[level],
            textColor=colors.HexColor(profile.accent if level == 1 else profile.text),
            spaceBefore=10,
            spaceAfter=5,
            keepWithNext=True,
        )
        for level in range(1, 7)
    }
    quote = ParagraphStyle(
        "CleanQuote",
        parent=body,
        leftIndent=18,
        textColor=colors.HexColor(profile.muted),
        fontName="Helvetica-Oblique",
    )
    code = ParagraphStyle(
        "CleanCode",
        parent=body,
        fontName="Courier",
        fontSize=max(8, profile.body_size - 1.5),
        leading=max(10, profile.body_size),
        leftIndent=10,
        rightIndent=10,
        borderPadding=7,
        backColor=colors.HexColor("#F4F6FA"),
    )
    story: list = [Paragraph(html.escape(parsed.title), heading_styles[1]), Spacer(1, 5)]
    pending_list: list[Block] = []

    def flush_list() -> None:
        nonlocal pending_list
        if not pending_list:
            return
        counters: dict[int, int] = {}
        for block in pending_list:
            for level in tuple(counters):
                if level > block.level:
                    counters.pop(level)
            if block.ordered:
                counters[block.level] = counters.get(block.level, 0) + 1
                marker = f"{counters[block.level]}."
            else:
                marker = "•"
                counters.pop(block.level, None)
            item_style = ParagraphStyle(
                f"ListItem-{block.level}-{int(block.ordered)}",
                parent=body,
                leftIndent=20 + block.level * 20,
                firstLineIndent=0,
                bulletIndent=6 + block.level * 20,
                spaceAfter=3,
            )
            story.append(Paragraph(_pdf_spans(block.spans), item_style, bulletText=marker))
        story.append(Spacer(1, 4))
        pending_list = []

    for block in _display_blocks(parsed):
        if block.kind == "list_item":
            if pending_list and pending_list[-1].ordered != block.ordered:
                flush_list()
            pending_list.append(block)
            continue
        flush_list()
        if block.kind == "heading":
            story.append(
                Paragraph(_pdf_spans(block.spans), heading_styles[min(max(block.level, 1), 6)])
            )
        elif block.kind == "paragraph":
            story.append(Paragraph(_pdf_spans(block.spans), body))
        elif block.kind == "blockquote":
            story.append(Paragraph(_pdf_spans(block.spans), quote))
        elif block.kind == "code":
            story.append(Preformatted(block.text, code, maxLineLength=95))
        elif block.kind == "horizontal_rule":
            story.append(
                HRFlowable(width="100%", thickness=0.7, color=colors.HexColor(profile.accent))
            )
        elif block.kind == "table" and block.rows:
            table_data = [
                [Paragraph(html.escape(cell), body) for cell in row] for row in block.rows
            ]
            column_count = max(len(row) for row in block.rows)
            for row in table_data:
                row.extend([Paragraph("", body)] * (column_count - len(row)))
            available_width = page_size[0] - 2 * profile.margin_inches * inch
            table = Table(
                table_data, colWidths=[available_width / column_count] * column_count, repeatRows=1
            )
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EAF0FB")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor(profile.text)),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#C9D2E3")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                        ("TOPPADDING", (0, 0), (-1, -1), 5),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ]
                )
            )
            story.append(KeepTogether([table, Spacer(1, 8)]))
    flush_list()
    document.build(story)
    return _artifact("pdf", path)


def to_plain_text(parsed: ParsedDocument) -> str:
    lines: list[str] = []
    for block in parsed.blocks:
        if block.kind == "heading":
            lines.extend(
                [block.text, "=" * min(72, len(block.text))] if block.level == 1 else [block.text]
            )
        elif block.kind == "list_item":
            marker = "1." if block.ordered else "-"
            lines.append(f"{'  ' * block.level}{marker} {block.text}")
        elif block.kind == "blockquote":
            lines.append(f"> {block.text}")
        elif block.kind == "code":
            lines.append(block.text)
        elif block.kind == "table":
            widths = [
                max(len(row[i]) if i < len(row) else 0 for row in block.rows)
                for i in range(max(map(len, block.rows)))
            ]
            for row in block.rows:
                lines.append(
                    " | ".join(
                        (row[i] if i < len(row) else "").ljust(widths[i])
                        for i in range(len(widths))
                    ).rstrip()
                )
        elif block.kind == "horizontal_rule":
            lines.append("-" * 48)
        else:
            lines.append(block.text)
        lines.append("")
    return "\n".join(lines).rstrip() + "\n"


def to_markdown(parsed: ParsedDocument) -> str:
    lines: list[str] = []
    for block in parsed.blocks:
        if block.kind == "heading":
            lines.append(f"{'#' * min(max(block.level, 1), 6)} {block.text}")
        elif block.kind == "list_item":
            marker = "1." if block.ordered else "-"
            lines.append(f"{'  ' * block.level}{marker} {block.text}")
        elif block.kind == "blockquote":
            lines.append(f"> {block.text}")
        elif block.kind == "code":
            lines.extend([f"```{block.language or ''}", block.text, "```"])
        elif block.kind == "table" and block.rows:
            header = block.rows[0]
            lines.append("| " + " | ".join(header) + " |")
            lines.append("| " + " | ".join("---" for _ in header) + " |")
            for row in block.rows[1:]:
                lines.append("| " + " | ".join(row) + " |")
        elif block.kind == "horizontal_rule":
            lines.append("---")
        else:
            value = ""
            for span in block.spans:
                text = span.text
                if span.code:
                    text = f"`{text}`"
                if span.bold:
                    text = f"**{text}**"
                if span.italic:
                    text = f"*{text}*"
                if span.href:
                    text = f"[{text}]({span.href})"
                value += text
            lines.append(value)
        lines.append("")
    return "\n".join(lines).rstrip() + "\n"


def generate_text(parsed: ParsedDocument, path: Path) -> GeneratedArtifact:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(to_plain_text(parsed), encoding="utf-8")
    return _artifact("text", path)


def generate_markdown(parsed: ParsedDocument, path: Path) -> GeneratedArtifact:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(to_markdown(parsed), encoding="utf-8")
    return _artifact("markdown", path)


GENERATORS = {
    "docx": generate_docx,
    "pdf": generate_pdf,
    "text": generate_text,
    "markdown": generate_markdown,
}
